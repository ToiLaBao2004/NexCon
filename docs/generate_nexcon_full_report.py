# -*- coding: utf-8 -*-
from pathlib import Path

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "generated_report_assets"
OUT = DOCS / "Bao_cao_PTTK_HTTT_NexCon_Tong_hop.docx"
LOGO = ROOT / "frontend" / "public" / "NexCon.jpeg"

PRIMARY = "1F4E79"
HEADER_FILL = "1F4E79"
LIGHT_FILL = "EAF2F8"
GRAY_FILL = "F3F4F6"
FONT = "Times New Roman"


def set_run_font(run, size=13, bold=None, italic=None, color=None, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=10, align=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    lines = str(text).split("\n")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        set_run_font(run, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths=None, font_size=10, shade_header=True):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, color="FFFFFF" if shade_header else None, size=font_size)
        if shade_header:
            set_cell_shading(cell, HEADER_FILL)
        if widths:
            cell.width = widths[index]
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value, size=font_size)
            if widths:
                cells[index].width = widths[index]
    doc.add_paragraph()
    return table


def add_center(doc, text, size=13, bold=False, italic=False, color=None, space_after=6):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return paragraph


def add_para(doc, text="", bold=False, italic=False, align=None, first_line=True, size=13, space_after=6):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(space_after)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Inches(0.35)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading("", level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=16, bold=True, color=PRIMARY)
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(12)
    elif level == 2:
        set_run_font(run, size=14, bold=True, color=PRIMARY)
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=13, bold=True)
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        set_run_font(run, size=13)


def add_numbered(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(item)
        set_run_font(run, size=13)


def add_caption(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, italic=True)


def add_code(doc, text):
    for line in text.strip("\n").split("\n"):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(line)
        set_run_font(run, size=9, font="Courier New")
    doc.add_paragraph()


def add_use_case(doc, code, name, actors, goal, pre, post, flow, alt, note=""):
    add_heading(doc, f"{code}. {name}", 3)
    rows = [
        ("Tác nhân", actors),
        ("Mục tiêu", goal),
        ("Điều kiện tiên quyết", pre),
        ("Điều kiện sau", post),
        ("Luồng sự kiện chính", flow),
        ("Luồng thay thế / ngoại lệ", alt),
    ]
    if note:
        rows.append(("Ghi chú", note))
    add_table(doc, ["Mục", "Nội dung"], rows, widths=[Inches(1.8), Inches(5.7)], font_size=10)


def add_sequence(doc, title, rows):
    add_heading(doc, title, 2)
    add_table(
        doc,
        ["STT", "Đối tượng gửi", "Đối tượng nhận", "Thông điệp / xử lý"],
        rows,
        widths=[Inches(0.5), Inches(1.5), Inches(1.6), Inches(4.2)],
        font_size=9,
    )


def add_test_table(doc, rows):
    add_table(
        doc,
        ["Mã", "Chức năng", "Điều kiện", "Bước kiểm thử", "Kết quả mong đợi"],
        rows,
        widths=[Inches(0.5), Inches(1.35), Inches(1.35), Inches(2.0), Inches(2.35)],
        font_size=8,
    )


def draw_box(ax, xy, text, width=1.8, height=0.55, color="#EAF2F8"):
    x, y = xy
    rect = plt.Rectangle((x, y), width, height, facecolor=color, edgecolor="#1F4E79", linewidth=1.2)
    ax.add_patch(rect)
    ax.text(x + width / 2, y + height / 2, text, ha="center", va="center", fontsize=8.5, wrap=True)


def draw_arrow(ax, start, end):
    ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", color="#333333", lw=1.1))


def create_diagrams():
    ASSETS.mkdir(parents=True, exist_ok=True)
    plt.rcParams["font.family"] = "DejaVu Sans"

    architecture = ASSETS / "architecture.png"
    fig, ax = plt.subplots(figsize=(10, 5.8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    draw_box(ax, (0.3, 4.6), "Web Browser\nAndroid App", 1.8)
    draw_box(ax, (2.6, 4.6), "React + Vite\nCapacitor", 1.8)
    draw_box(ax, (5.0, 4.6), "Express API\nSocket.IO", 1.8)
    draw_box(ax, (7.5, 4.6), "LiveKit\nWebRTC", 1.7)
    draw_box(ax, (0.7, 2.6), "MongoDB\nDữ liệu chính", 1.7, color="#F3F4F6")
    draw_box(ax, (2.8, 2.6), "Redis\nPresence/Call/Queue", 1.9, color="#F3F4F6")
    draw_box(ax, (5.1, 2.6), "Cloudinary\nMedia", 1.7, color="#F3F4F6")
    draw_box(ax, (7.1, 2.6), "Gemini +\nAssemblyAI", 1.8, color="#F3F4F6")
    draw_box(ax, (1.5, 0.8), "BullMQ Workers\nReminder/Cleanup/Expiry", 2.2, color="#FFF2CC")
    draw_box(ax, (4.4, 0.8), "Web Push + FCM\nEmail", 2.0, color="#FFF2CC")
    draw_box(ax, (7.0, 0.8), "Admin Dashboard\nObservability", 2.0, color="#FFF2CC")
    for s, e in [
        ((2.1, 4.9), (2.6, 4.9)),
        ((4.4, 4.9), (5.0, 4.9)),
        ((6.8, 4.9), (7.5, 4.9)),
        ((5.9, 4.6), (1.55, 3.15)),
        ((5.9, 4.6), (3.75, 3.15)),
        ((5.9, 4.6), (5.95, 3.15)),
        ((5.9, 4.6), (8.0, 3.15)),
        ((3.75, 2.6), (2.6, 1.35)),
        ((5.0, 4.6), (5.4, 1.35)),
        ((5.0, 4.6), (8.0, 1.35)),
    ]:
        draw_arrow(ax, s, e)
    ax.set_title("Kiến trúc tổng thể NexCon", fontsize=13, fontweight="bold")
    fig.tight_layout()
    fig.savefig(architecture, dpi=180)
    plt.close(fig)

    usecase = ASSETS / "usecase_overview.png"
    fig, ax = plt.subplots(figsize=(10, 5.8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis("off")
    draw_box(ax, (0.3, 4.8), "Khách", 1.2, color="#F3F4F6")
    draw_box(ax, (0.3, 2.8), "Người dùng", 1.2, color="#F3F4F6")
    draw_box(ax, (0.3, 1.0), "Admin", 1.2, color="#F3F4F6")
    draw_box(ax, (8.4, 3.6), "AI\nModeration", 1.3, color="#F3F4F6")
    cases = [
        ("Đăng ký/Đăng nhập", 2.5, 4.8),
        ("Kết bạn/Chặn", 2.4, 3.5),
        ("Chat realtime", 4.6, 3.5),
        ("Gọi/Họp", 6.8, 3.5),
        ("Reminder/Thông báo", 2.4, 2.2),
        ("Report vi phạm", 4.6, 2.2),
        ("Quản trị user", 4.6, 1.0),
        ("Dashboard/AI review", 6.8, 1.0),
    ]
    for label, x, y in cases:
        draw_box(ax, (x, y), label, 1.8, color="#EAF2F8")
    arrows = [
        ((1.5, 5.05), (2.5, 5.05)),
        ((1.5, 3.1), (2.4, 3.75)),
        ((1.5, 3.1), (4.6, 3.75)),
        ((1.5, 3.1), (6.8, 3.75)),
        ((1.5, 3.1), (2.4, 2.5)),
        ((1.5, 3.1), (4.6, 2.5)),
        ((1.5, 1.3), (4.6, 1.3)),
        ((1.5, 1.3), (6.8, 1.3)),
        ((6.4, 2.5), (8.4, 3.9)),
        ((8.1, 1.3), (8.8, 3.6)),
    ]
    for s, e in arrows:
        draw_arrow(ax, s, e)
    ax.set_title("Use Case tổng quát NexCon", fontsize=13, fontweight="bold")
    fig.tight_layout()
    fig.savefig(usecase, dpi=180)
    plt.close(fig)

    pipeline = ASSETS / "ai_moderation_pipeline.png"
    fig, ax = plt.subplots(figsize=(10, 4.8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")
    steps = [
        ("Người gửi\nsoạn nội dung", 0.3, 2.8),
        ("Message API\nkiểm quyền", 2.0, 2.8),
        ("Lưu + emit\npending_review", 3.8, 2.8),
        ("Local signal +\nGemini", 5.8, 3.5),
        ("AssemblyAI\nVoice transcript", 5.8, 1.7),
        ("Decision\napproved/rejected/skipped", 7.8, 2.8),
        ("Cập nhật message\nẨn/cleanup nếu vi phạm", 7.8, 0.8),
    ]
    for label, x, y in steps:
        draw_box(ax, (x, y), label, 1.55, color="#EAF2F8" if "Decision" not in label else "#FFF2CC")
    for s, e in [
        ((1.85, 3.05), (2.0, 3.05)),
        ((3.55, 3.05), (3.8, 3.05)),
        ((5.35, 3.05), (5.8, 3.75)),
        ((5.35, 3.05), (5.8, 1.95)),
        ((7.35, 3.75), (7.8, 3.05)),
        ((7.35, 1.95), (7.8, 3.05)),
        ((8.55, 2.8), (8.55, 1.35)),
    ]:
        draw_arrow(ax, s, e)
    ax.text(8.6, 2.55, "confidence >= ngưỡng\n=> ẩn nội dung đã gửi", fontsize=8, ha="center")
    ax.set_title("Pipeline AI kiểm duyệt nội dung", fontsize=13, fontweight="bold")
    fig.tight_layout()
    fig.savefig(pipeline, dpi=180)
    plt.close(fig)

    data_model = ASSETS / "data_model.png"
    fig, ax = plt.subplots(figsize=(10, 6.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6.5)
    ax.axis("off")
    boxes = [
        ("User", 0.5, 5.3),
        ("Session\nUserStatus", 0.5, 4.1),
        ("Friend\nFriendRequest\nBlockUser", 0.5, 2.8),
        ("Conversation", 3.4, 5.3),
        ("Message", 3.4, 4.0),
        ("Meeting", 6.5, 5.3),
        ("Reminder", 6.5, 4.0),
        ("Report", 3.4, 2.4),
        ("Notification", 6.5, 2.4),
        ("LockAppeal\nAuditLog\nPushSubscription\nOtp", 3.4, 0.9),
    ]
    for label, x, y in boxes:
        draw_box(ax, (x, y), label, 2.0, 0.75, color="#EAF2F8")
    for s, e in [
        ((2.5, 5.65), (3.4, 5.65)),
        ((4.4, 5.3), (4.4, 4.75)),
        ((5.4, 5.65), (6.5, 5.65)),
        ((5.4, 4.35), (6.5, 4.35)),
        ((2.5, 5.4), (3.4, 2.75)),
        ((4.4, 4.0), (4.4, 3.15)),
        ((2.5, 4.45), (3.4, 1.3)),
        ((2.5, 3.15), (3.4, 5.55)),
        ((5.4, 2.75), (6.5, 2.75)),
        ((8.5, 4.35), (8.5, 2.95)),
    ]:
        draw_arrow(ax, s, e)
    ax.set_title("Data Model / ERD rút gọn", fontsize=13, fontweight="bold")
    fig.tight_layout()
    fig.savefig(data_model, dpi=180)
    plt.close(fig)

    return {
        "architecture": architecture,
        "usecase": usecase,
        "pipeline": pipeline,
        "data_model": data_model,
    }


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(0.8)
    for style_name in ["Normal", "Body Text", "List Bullet", "List Number"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = FONT
            style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
            style.font.size = Pt(13)
    return doc


def add_front_matter(doc):
    add_center(doc, "TRƯỜNG ĐẠI HỌC SƯ PHẠM KỸ THUẬT TP. HỒ CHÍ MINH", 14, True)
    add_center(doc, "KHOA CÔNG NGHỆ THÔNG TIN", 14, True)
    add_center(doc, "-----o0o-----", 13)
    if LOGO.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(LOGO), width=Inches(1.3))
    add_center(doc, "BÁO CÁO PHÂN TÍCH THIẾT KẾ HỆ THỐNG THÔNG TIN", 17, True, color=PRIMARY, space_after=16)
    add_center(doc, "TÊN ĐỀ TÀI:", 14, True)
    add_center(doc, "PHÂN TÍCH THIẾT KẾ HỆ THỐNG ỨNG DỤNG GIAO TIẾP REALTIME NEXCON", 15, True, color=PRIMARY)
    add_center(doc, "TÍCH HỢP CHAT, GỌI AUDIO/VIDEO VÀ AI KIỂM DUYỆT NỘI DUNG", 15, True, color=PRIMARY, space_after=28)
    add_center(doc, "GVHD: ............................................................", 13)
    add_center(doc, "Lớp HP: ............................................................", 13)
    add_center(doc, "Nhóm thực hiện: ..................................................", 13)
    add_center(doc, "Học kỳ: ......", 13)
    add_center(doc, "Năm học: 2025 - 2026", 13)
    add_center(doc, "Thành phố Hồ Chí Minh, tháng 6 năm 2026", 13)
    doc.add_page_break()

    add_center(doc, "DANH SÁCH SINH VIÊN NHÓM THỰC HIỆN", 15, True, color=PRIMARY)
    add_table(
        doc,
        ["STT", "MSSV", "Họ và tên", "Vai trò / ghi chú"],
        [
            ["1", "", "", "Nhóm trưởng"],
            ["2", "", "", "Thành viên"],
            ["3", "", "", "Thành viên"],
            ["4", "", "", "Thành viên"],
        ],
        widths=[Inches(0.6), Inches(1.4), Inches(3.2), Inches(2.0)],
        font_size=11,
    )
    add_center(doc, "NHẬN XÉT CỦA GIẢNG VIÊN", 15, True, color=PRIMARY)
    for _ in range(10):
        add_para(doc, "................................................................................................................................................................", first_line=False, space_after=2)
    add_para(doc, "Tp. Hồ Chí Minh, ngày ...... tháng ...... năm 2026", first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_para(doc, "Giảng viên chấm điểm", first_line=False, align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()

    add_heading(doc, "LỜI CẢM ƠN", 1)
    add_para(doc, "Nhóm xin gửi lời cảm ơn chân thành đến quý thầy cô đã hướng dẫn, góp ý và tạo điều kiện để nhóm thực hiện đề tài phân tích thiết kế hệ thống thông tin. Quá trình thực hiện báo cáo giúp nhóm hệ thống hóa kiến thức về khảo sát yêu cầu, mô hình hóa UML, thiết kế dữ liệu, kiến trúc phần mềm realtime và kiểm thử hệ thống.")
    add_para(doc, "Đề tài NexCon được xây dựng trong bối cảnh nhu cầu giao tiếp trực tuyến, làm việc từ xa, họp trực tuyến và quản trị an toàn cộng đồng ngày càng quan trọng. Nhóm đã tham khảo bố cục từ các báo cáo mẫu được cung cấp, đồng thời điều chỉnh nội dung theo chức năng thực tế của dự án.")
    add_para(doc, "Báo cáo khó tránh khỏi thiếu sót trong quá trình biên soạn. Nhóm mong nhận được nhận xét từ giảng viên để hoàn thiện hơn cả về tài liệu và sản phẩm.")
    doc.add_page_break()

    add_heading(doc, "TÓM TẮT ĐỀ TÀI", 1)
    add_para(doc, "NexCon là nền tảng giao tiếp realtime đa nền tảng, hỗ trợ web và Android. Hệ thống cung cấp các chức năng cốt lõi như xác thực tài khoản, quản lý hồ sơ, kết bạn, chat cá nhân, chat nhóm, gửi media, gọi audio/video, phòng họp LiveKit, reminder, thông báo đa kênh, report vi phạm, admin dashboard và AI moderation.")
    add_para(doc, "Điểm nhấn của đề tài là tích hợp AI để kiểm duyệt nội dung theo tiêu chuẩn cộng đồng bằng cơ chế hậu kiểm. Tin nhắn được lưu và phát realtime trước với trạng thái pending_review; sau đó hệ thống kết hợp local signal, Google Gemini và AssemblyAI để phân tích text, link, image, file metadata và transcript của voice message. Khi phát hiện nội dung vi phạm với độ tin cậy đủ cao, hệ thống ẩn nội dung đã gửi, đánh dấu reportStatus, cleanup media nếu cần, ghi nhận vi phạm, cảnh báo người dùng hoặc khóa tài khoản khi vượt ngưỡng.")
    add_para(doc, "Về kiến trúc, NexCon kết hợp React, Vite, TypeScript, Node.js, Express, MongoDB, Redis, Socket.IO, BullMQ, LiveKit, Cloudinary, Web Push, Firebase Cloud Messaging và Capacitor. Báo cáo trình bày đầy đủ phần khảo sát, yêu cầu chức năng/phi chức năng, use case, activity, sequence, class diagram, data model/ERD, AI moderation pipeline, công nghệ, cài đặt, kiểm thử, giao diện và hướng phát triển.")
    doc.add_page_break()

    add_heading(doc, "MỤC LỤC", 1)
    toc_items = [
        "CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI",
        "CHƯƠNG 2. KHẢO SÁT HIỆN TRẠNG VÀ XÁC ĐỊNH YÊU CẦU",
        "CHƯƠNG 3. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ",
        "CHƯƠNG 4. PHÂN TÍCH THIẾT KẾ HỆ THỐNG",
        "CHƯƠNG 5. CÀI ĐẶT VÀ KIỂM THỬ",
        "CHƯƠNG 6. THIẾT KẾ GIAO DIỆN",
        "CHƯƠNG 7. KẾT LUẬN",
        "TÀI LIỆU THAM KHẢO",
        "PHỤ LỤC",
    ]
    for item in toc_items:
        add_para(doc, item, bold=True, first_line=False, space_after=2)
    add_heading(doc, "DANH MỤC HÌNH ẢNH", 1)
    for item in [
        "Hình 4.1. Kiến trúc tổng thể NexCon",
        "Hình 4.2. Use Case tổng quát NexCon",
        "Hình 4.3. Pipeline AI kiểm duyệt nội dung",
        "Hình 4.4. Data Model / ERD rút gọn",
    ]:
        add_para(doc, item, first_line=False, space_after=2)
    add_heading(doc, "DANH MỤC BẢNG BIỂU", 1)
    for item in [
        "Bảng 2.1. Đối tượng sử dụng hệ thống",
        "Bảng 2.2. Yêu cầu chức năng",
        "Bảng 2.3. Yêu cầu phi chức năng",
        "Bảng 4.1. Đặc tả Use Case",
        "Bảng 4.2. Class Diagram dạng bảng",
        "Bảng 4.3. Thiết kế Data Model",
        "Bảng 5.1. Cấu hình CI/CD và deploy production",
        "Bảng 5.2. Xử lý vận hành backend multi-replica 6 replicas",
        "Bảng 5.3. Test cases và kết quả mong đợi",
        "Bảng 6.1. Danh sách màn hình giao diện",
    ]:
        add_para(doc, item, first_line=False, space_after=2)
    doc.add_page_break()


def add_chapter_1(doc):
    add_heading(doc, "CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Tính cấp thiết của đề tài", 2)
    add_para(doc, "Giao tiếp trực tuyến đã trở thành một phần không thể thiếu trong học tập, làm việc và sinh hoạt hằng ngày. Người dùng không chỉ cần nhắn tin mà còn cần gọi audio/video, họp nhóm, chia sẻ file, nhận thông báo tức thời và quản lý lịch nhắc hẹn trong cùng một nền tảng.")
    add_para(doc, "Tuy nhiên, các ứng dụng giao tiếp hiện đại cũng đối mặt với các vấn đề về spam, quấy rối, nội dung độc hại, link không an toàn, hình ảnh vi phạm, lạm dụng tài khoản và khó khăn trong việc quản trị cộng đồng. Nếu thiếu cơ chế kiểm duyệt và báo cáo minh bạch, trải nghiệm người dùng sẽ bị ảnh hưởng nghiêm trọng.")
    add_para(doc, "NexCon được đề xuất như một hệ thống giao tiếp realtime tích hợp AI moderation. Mục tiêu là xây dựng một nền tảng vừa có khả năng trao đổi nhanh, ổn định, đa thiết bị, vừa có cơ chế phát hiện và xử lý nội dung vi phạm tiêu chuẩn cộng đồng.")

    add_heading(doc, "1.2. Mục tiêu đề tài", 2)
    add_bullets(doc, [
        "Phân tích và thiết kế hệ thống chat/call realtime phục vụ web và Android.",
        "Xây dựng mô hình chức năng cho xác thực, hồ sơ, bạn bè, hội thoại, chat, call, meeting, reminder, notification, report và admin.",
        "Thiết kế cơ chế AI moderation hậu kiểm cho text, link, image, file metadata và voice transcript.",
        "Thiết kế data model phù hợp với MongoDB, Redis và các tác vụ realtime.",
        "Đề xuất kiến trúc có khả năng mở rộng, bảo mật phiên đăng nhập, quản lý media và vận hành worker nền.",
        "Xây dựng bộ test case chức năng để đánh giá các luồng chính của hệ thống.",
    ])

    add_heading(doc, "1.3. Đối tượng và phạm vi nghiên cứu", 2)
    add_table(
        doc,
        ["Nội dung", "Phạm vi"],
        [
            ["Đối tượng sử dụng", "Khách, người dùng đã đăng nhập, quản trị viên, AI moderation service và worker nền."],
            ["Nền tảng", "Web app responsive và Android app thông qua Capacitor. Repository hiện chưa có target iOS."],
            ["Chức năng chính", "Chat cá nhân/nhóm, gọi audio/video, phòng họp, reminder, notification, report, moderation, admin dashboard."],
            ["Dữ liệu", "User, session, friend, conversation, message, meeting, reminder, notification, report, violation, appeal và audit log."],
            ["Ngoài phạm vi", "Mã hóa end-to-end hoàn chỉnh, thanh toán, hệ thống mạng xã hội feed/post/reel/story và target iOS native."],
        ],
        widths=[Inches(1.8), Inches(5.7)],
    )

    add_heading(doc, "1.4. Phương pháp tiếp cận", 2)
    add_numbered(doc, [
        "Khảo sát các nền tảng giao tiếp và mạng xã hội phổ biến để xác định nhu cầu chức năng.",
        "Phân tích mã nguồn dự án NexCon để xác định thực thể dữ liệu, API, route, worker và realtime event đang có.",
        "Mô hình hóa yêu cầu bằng use case, activity diagram, sequence diagram, class diagram và data model.",
        "Thiết kế pipeline AI moderation hậu kiểm dựa trên local signal, Gemini, AssemblyAI, report và admin review.",
        "Xây dựng bảng test case theo từng nhóm chức năng để kiểm tra đầu ra mong đợi.",
    ])

    add_heading(doc, "1.5. Kết quả dự kiến đạt được", 2)
    add_bullets(doc, [
        "Một tài liệu phân tích thiết kế đầy đủ cho hệ thống NexCon.",
        "Bộ mô hình chức năng, hành vi và dữ liệu phục vụ triển khai hoặc bảo trì hệ thống.",
        "Định nghĩa rõ trách nhiệm của client, API, Socket.IO, worker, Redis, MongoDB và dịch vụ AI.",
        "Danh sách test case hỗ trợ kiểm thử chức năng và kiểm thử hồi quy.",
        "Định hướng phát triển tiếp theo về bảo mật, AI moderation, scale realtime và mobile.",
    ])
    doc.add_page_break()


def add_chapter_2(doc):
    add_heading(doc, "CHƯƠNG 2. KHẢO SÁT HIỆN TRẠNG VÀ XÁC ĐỊNH YÊU CẦU", 1)
    add_heading(doc, "2.1. Bối cảnh sản phẩm", 2)
    add_para(doc, "NexCon hướng đến nhóm người dùng cần một nền tảng giao tiếp nhanh, trực quan và có mức độ kiểm soát cộng đồng tốt. Khác với website thương mại điện tử tập trung vào sản phẩm, giỏ hàng và đơn hàng, NexCon tập trung vào quan hệ người dùng, hội thoại, trạng thái realtime, cuộc gọi, thông báo và quản trị nội dung.")
    add_para(doc, "Các yêu cầu quan trọng của hệ thống realtime gồm độ trễ thấp, đồng bộ nhiều thiết bị, xử lý offline/online, bảo vệ session, kiểm tra quyền truy cập từng hội thoại, kiểm soát media và phát hiện nội dung vi phạm.")

    add_heading(doc, "2.2. Khảo sát một số hệ thống tương tự", 2)
    add_table(
        doc,
        ["Hệ thống", "Điểm mạnh", "Hạn chế / cơ hội cải thiện cho NexCon"],
        [
            ["Messenger", "Chat realtime, media, call, group chat, notification mạnh.", "Kiểm soát dữ liệu và tùy biến quy trình admin/report phụ thuộc nền tảng."],
            ["Zalo", "Phổ biến tại Việt Nam, nhắn tin/gọi ổn định, hỗ trợ nhóm.", "Tích hợp AI moderation hoặc dashboard vận hành không phải trọng tâm với người dùng cuối."],
            ["Discord", "Mạnh về cộng đồng, voice channel, role và nhóm.", "Giao diện nhiều tính năng có thể phức tạp với nhóm người dùng phổ thông."],
            ["Slack", "Tốt cho làm việc nhóm, channel, search, notification.", "Không tối ưu cho social friend graph và Android consumer call experience."],
        ],
        widths=[Inches(1.3), Inches(3.0), Inches(3.2)],
    )

    add_heading(doc, "2.3. Đặc tả bài toán", 2)
    add_para(doc, "Bài toán đặt ra là thiết kế một hệ thống giao tiếp realtime cho phép người dùng tạo tài khoản, kết nối bạn bè, nhắn tin cá nhân hoặc nhóm, gọi audio/video, tổ chức phòng họp, tạo reminder và nhận thông báo đa kênh. Hệ thống phải có cơ chế phát hiện, báo cáo và xử lý nội dung vi phạm tiêu chuẩn cộng đồng.")
    add_para(doc, "Đối với backend, bài toán không chỉ là lưu tin nhắn mà còn phải đảm bảo tính nhất quán của unread count, delivery/read status, trạng thái call, permission theo hội thoại, block/lock user, media signed URL, worker hết hạn dữ liệu, và audit log cho quản trị.")

    add_heading(doc, "2.4. Đối tượng sử dụng", 2)
    add_table(
        doc,
        ["Vai trò", "Mô tả", "Chức năng chính"],
        [
            ["Khách", "Người chưa đăng nhập hoặc chưa có tài khoản.", "Đăng ký, đăng nhập, OTP, quên mật khẩu, Google OAuth, gửi kháng cáo khi tài khoản bị khóa."],
            ["Người dùng", "Người đã xác thực và sử dụng chức năng giao tiếp.", "Kết bạn, chặn, cập nhật hồ sơ, chat, gọi, họp, reminder, notification, report, xem trạng thái moderation."],
            ["Quản trị viên", "Người có quyền giám sát và xử lý an toàn cộng đồng.", "Dashboard, observability, user management, report review, AI review, lock/unlock, appeal review, audit log."],
            ["AI moderation service", "Dịch vụ hỗ trợ phân tích nội dung tự động.", "Hậu kiểm text, link, image, file metadata, transcript; trả category, confidence, reason và trạng thái approved/rejected/skipped."],
            ["Worker nền", "Tiến trình xử lý tác vụ không đồng bộ.", "Reminder, disappearing expiry, cleanup group/conversation, realtime timeout."],
        ],
        widths=[Inches(1.35), Inches(2.35), Inches(4.0)],
    )

    add_heading(doc, "2.5. Yêu cầu chức năng", 2)
    add_table(
        doc,
        ["Nhóm", "Yêu cầu chức năng"],
        [
            ["Xác thực", "Đăng ký, đăng nhập, Google OAuth, OTP, reset password, refresh token, signout một phiên/toàn bộ phiên, quản lý thiết bị."],
            ["Hồ sơ", "Cập nhật displayName, avatar, bio, phone, profileVisibility, bài hát hồ sơ, trạng thái online/manual/auto."],
            ["Bạn bè", "Tìm user, gửi/hủy/chấp nhận/từ chối/resend lời mời, gợi ý bạn bè, nickname, unfriend, block/unblock."],
            ["Chat", "Direct/group conversation, text/image/audio/file/link/sticker/system message, reply, reaction, recall, pin, forward, mention, search, media sidebar."],
            ["Nhóm", "Tạo nhóm, đổi tên/avatar, thêm/xóa thành viên, chuyển admin, approval queue, rời nhóm, giải tán nhóm, cleanup dữ liệu."],
            ["Disappearing", "Bật/tắt chế độ tự xóa, TTL 24 giờ cho tin mới, soft-delete, screenshot detection Android 14+."],
            ["Call/Meeting", "Direct call, group call, LiveKit meeting room, waiting room, admit/reject, scheduled meeting, call history system message."],
            ["Reminder", "Reminder cá nhân/chung, tạo từ tin nhắn hoặc lịch họp, repeat, snooze, dismiss, participation."],
            ["Notification", "In-app notification, Web Push, FCM, local notification Android, mark read/unread/delete."],
            ["Moderation/Report", "Hậu kiểm text/link/image/file metadata/transcript, report message/user, snapshot evidence, AI review, violation history, auto lock, appeal."],
            ["Admin", "Stats, observability, users, profile, messages, conversations, audit logs, reports, appeals, lock/unlock user."],
        ],
        widths=[Inches(1.5), Inches(6.0)],
        font_size=9,
    )

    add_heading(doc, "2.6. Yêu cầu phi chức năng", 2)
    add_table(
        doc,
        ["Thuộc tính", "Yêu cầu"],
        [
            ["Hiệu năng", "Tin nhắn và typing event phải phản hồi nhanh; truy vấn hội thoại dùng index theo user/conversation/createdAt; global search hỗ trợ stream."],
            ["Độ tin cậy", "Redis lưu call state và BullMQ job; worker có thể reload pending reminder/cleanup sau khi server khởi động lại."],
            ["Bảo mật", "JWT, refresh token theo session, cookie, rate limit, auth middleware, role middleware, signed media URL, revoke session khi lock."],
            ["Riêng tư", "Profile visibility, block user, mute, invisible status, media authenticated, chỉ trả presence phù hợp quan hệ."],
            ["Mở rộng", "Socket.IO Redis adapter cho nhiều backend replica; Redis pub/sub đồng bộ room/user event."],
            ["Bảo trì", "Backend tách controller, route, middleware, service, model, worker; frontend tách pages, stores, services, components."],
            ["Khả dụng mobile", "Capacitor hỗ trợ Android, FCM, local notification, incoming call và native screenshot bridge."],
            ["An toàn cộng đồng", "Report, AI moderation, admin review, violation decay, auto lock và appeal minh bạch."],
        ],
        widths=[Inches(1.6), Inches(5.9)],
        font_size=9,
    )

    add_heading(doc, "2.7. Dữ liệu cần lưu trữ", 2)
    add_table(
        doc,
        ["Nhóm dữ liệu", "Thực thể", "Mô tả"],
        [
            ["Tài khoản", "User, Session, Otp, UserStatus", "Thông tin đăng nhập, hồ sơ, role, lock, session, OTP và trạng thái online."],
            ["Quan hệ", "Friend, FriendRequest, BlockUser", "Quan hệ bạn bè, lời mời và danh sách chặn."],
            ["Hội thoại", "Conversation, Message", "Direct/group conversation, participant, group admin, unread, message, media, reaction, mention, reportStatus."],
            ["Gọi/họp", "Meeting, Redis call state", "Room, host, waiting room, participants và state cuộc gọi realtime."],
            ["Nhắc hẹn", "Reminder", "Reminder cá nhân/chung, thời gian nhắc, repeat, snooze, participation, notify channel."],
            ["Thông báo", "Notification, PushSubscription", "In-app notification, Web Push subscription và FCM token."],
            ["An toàn cộng đồng", "Report, LockAppeal, User.moderation", "Report, evidence snapshot, AI result, violation counter, lock appeal."],
            ["Vận hành", "AuditLog, system metrics", "Request log, duration, status code, observability và dashboard admin."],
        ],
        widths=[Inches(1.5), Inches(2.2), Inches(3.8)],
        font_size=9,
    )

    add_heading(doc, "2.8. Quy trình nghiệp vụ chính", 2)
    add_bullets(doc, [
        "Quy trình xác thực: đăng ký/đăng nhập, OTP hoặc Google, tạo session, cấp token, Socket.IO xác thực và join user room.",
        "Quy trình kết bạn: tìm kiếm, gửi request, nhận realtime notification, chấp nhận/từ chối, tạo Friend và cho phép direct conversation.",
        "Quy trình chat: kiểm quyền, kiểm block/lock, upload media, lưu message với pending_review nếu cần, cập nhật conversation, phát socket realtime, sau đó chạy moderation nền và lên lịch notification khi không bị chặn.",
        "Quy trình call: signaling qua Socket.IO, kiểm friendship/membership, lưu Redis call state, cấp LiveKit token, phát event call lifecycle.",
        "Quy trình moderation/report: AI hậu kiểm nội dung đã gửi hoặc nội dung bị report, cập nhật approved/skipped/rejected, ẩn nội dung vi phạm, admin review, violation history, cảnh báo hoặc khóa tài khoản, appeal khi bị khóa.",
        "Quy trình worker: reminder trigger, disappearing expiry, cleanup dữ liệu nhóm/hội thoại và realtime timeout.",
    ])
    doc.add_page_break()


def add_chapter_3(doc):
    add_heading(doc, "CHƯƠNG 3. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ", 1)
    sections = [
        ("3.1. Ứng dụng realtime và Socket.IO", "Ứng dụng realtime cần truyền sự kiện hai chiều giữa client và server với độ trễ thấp. Socket.IO hỗ trợ WebSocket fallback, room, namespace, reconnect và event tùy biến. Trong NexCon, Socket.IO được dùng cho presence, typing, message delivery, read status, notification, reminder, call signaling, group call và waiting room."),
        ("3.2. WebRTC và LiveKit", "WebRTC cho phép truyền audio/video trực tiếp giữa client. LiveKit cung cấp server và SDK để quản lý room, token, participant và media track. NexCon dùng backend để kiểm quyền và cấp token, còn LiveKit xử lý media stream."),
        ("3.3. MongoDB và Mongoose", "MongoDB phù hợp với dữ liệu linh hoạt như message metadata, reactions, mentions, group settings và moderation result. Mongoose giúp định nghĩa schema, validate, index, TTL và middleware pre-save."),
        ("3.4. Redis và BullMQ", "Redis được dùng cho presence, call state, Socket.IO adapter, violation counter và queue. BullMQ hỗ trợ job nền như reminder, disappearing expiry, cleanup group và realtime timeout."),
        ("3.5. AI Moderation", "AI moderation là quá trình phân tích nội dung nhằm phát hiện vi phạm tiêu chuẩn cộng đồng. NexCon sử dụng cơ chế hậu kiểm: message được lưu và emit trước ở trạng thái pending_review, sau đó hệ thống kết hợp local signal, Google Gemini cho text/link/image/file metadata và AssemblyAI cho audio transcript. Kết quả AI được dùng để cập nhật approved/rejected/skipped, ẩn nội dung vi phạm và hỗ trợ admin review."),
        ("3.6. Bảo mật phiên và quyền truy cập", "Hệ thống sử dụng JWT access token, refresh token theo Session, auth middleware, role middleware, rate limit, revoke session, signed media URL và kiểm tra membership trước mỗi thao tác message/conversation."),
        ("3.7. Web Push, FCM và notification", "Thông báo đa kênh giúp người dùng nhận message, mention, reminder, friend request, moderation status và incoming call khi không nhìn trực tiếp vào app. Web Push dùng cho trình duyệt, FCM dùng cho Android."),
        ("3.8. Capacitor cho Android", "Capacitor cho phép đóng gói React app thành Android app, đồng thời tích hợp FCM, local notification, Google Sign-In, full-screen incoming call và native screenshot callback."),
    ]
    for title, body in sections:
        add_heading(doc, title, 2)
        add_para(doc, body)

    add_heading(doc, "3.9. Công nghệ sử dụng trong dự án", 2)
    add_table(
        doc,
        ["Lớp", "Công nghệ", "Vai trò"],
        [
            ["Frontend", "React 19, TypeScript, Vite 7, Tailwind CSS, Zustand", "Giao diện web, quản lý state, gọi API, render realtime UI."],
            ["Mobile", "Capacitor, FCM, Local Notifications", "Đóng gói Android, push notification và native bridge."],
            ["Backend", "Node.js, Express 5, Mongoose", "REST API, xác thực, nghiệp vụ, kết nối MongoDB."],
            ["Realtime", "Socket.IO, Redis Adapter", "Presence, chat events, call signaling, room sync nhiều replica."],
            ["Database", "MongoDB", "Lưu dữ liệu chính của hệ thống."],
            ["Cache/Queue", "Redis, BullMQ", "State tạm thời, worker job và queue nền."],
            ["Media", "Cloudinary", "Lưu avatar và media message, cấp signed URL."],
            ["Call", "LiveKit", "Audio/video/WebRTC room."],
            ["AI", "Google Gemini, AssemblyAI", "Moderation text/link/image và transcription voice."],
            ["Notification", "Web Push, Firebase Cloud Messaging, Email", "Thông báo đa kênh và OTP."],
            ["Kiểm thử", "Node test, Vitest", "Unit test backend/frontend."],
        ],
        widths=[Inches(1.3), Inches(2.6), Inches(3.6)],
        font_size=9,
    )
    doc.add_page_break()


def add_chapter_4(doc, diagrams):
    add_heading(doc, "CHƯƠNG 4. PHÂN TÍCH THIẾT KẾ HỆ THỐNG", 1)
    add_heading(doc, "4.1. Kiến trúc tổng thể", 2)
    doc.add_picture(str(diagrams["architecture"]), width=Inches(6.7))
    add_caption(doc, "Hình 4.1. Kiến trúc tổng thể NexCon")
    add_para(doc, "Client web/Android giao tiếp với backend thông qua REST API và Socket.IO. LiveKit xử lý media audio/video. Backend lưu dữ liệu chính ở MongoDB, lưu trạng thái realtime/queue ở Redis, lưu media ở Cloudinary, gọi Gemini/AssemblyAI để kiểm duyệt và sử dụng Web Push/FCM/email cho thông báo.")

    add_heading(doc, "4.2. Use Case Diagram", 2)
    doc.add_picture(str(diagrams["usecase"]), width=Inches(6.7))
    add_caption(doc, "Hình 4.2. Use Case tổng quát NexCon")
    add_para(doc, "Use case tổng quát gồm bốn nhóm tác nhân chính: khách, người dùng, quản trị viên và AI moderation service. Người dùng tương tác với hệ thống qua chat/call/reminder/report, trong khi admin xử lý báo cáo, giám sát hệ thống và quản lý tài khoản.")

    add_heading(doc, "4.3. Đặc tả Use Case", 2)
    use_cases = [
        ("UC01", "Đăng ký tài khoản", "Khách", "Tạo tài khoản NexCon mới.", "Email chưa tồn tại, thông tin hợp lệ.", "User được tạo và có thể đăng nhập.", "1. Khách nhập thông tin.\n2. Hệ thống kiểm tra định dạng.\n3. Gửi OTP xác minh.\n4. Khách nhập OTP.\n5. Backend tạo User.", "Email trùng, OTP sai hoặc hết hạn thì yêu cầu nhập lại."),
        ("UC02", "Đăng nhập", "Khách, Admin", "Xác thực tài khoản và tạo session.", "Tài khoản tồn tại, chưa bị khóa.", "Session/token được cấp, socket có thể kết nối.", "1. Nhập email/mật khẩu hoặc Google.\n2. Backend kiểm tra rate limit.\n3. Kiểm tra password/Google token.\n4. Kiểm tra lock.\n5. Tạo Session.", "Nếu tài khoản bị khóa, trả restriction và hướng dẫn appeal."),
        ("UC03", "Quản lý phiên đăng nhập", "Người dùng", "Xem và thu hồi session trên các thiết bị.", "Đã đăng nhập.", "Session bị thu hồi và socket thiết bị đó bị ngắt.", "1. Mở settings/sessions.\n2. Chọn thu hồi một phiên hoặc tất cả.\n3. Backend xóa Session.\n4. Emit session-revoked.", "Session đã hết hạn thì tự bị TTL xóa."),
        ("UC04", "Cập nhật hồ sơ", "Người dùng", "Cập nhật thông tin cá nhân và quyền riêng tư.", "Đã đăng nhập.", "User profile được cập nhật.", "1. Sửa displayName, bio, phone, avatar, music.\n2. Backend validate.\n3. Upload avatar nếu có.\n4. Lưu User và emit profile-updated.", "File avatar sai định dạng hoặc dữ liệu vượt giới hạn thì báo lỗi."),
        ("UC05", "Kết bạn", "Người dùng", "Tạo quan hệ bạn bè để chat/call trực tiếp.", "Hai user tồn tại, không bị chặn/khóa.", "Friend được tạo.", "1. Tìm user.\n2. Gửi request.\n3. Người nhận được notification.\n4. Người nhận accept.\n5. Backend tạo Friend.", "Reject/cancel/resend xử lý theo trạng thái request."),
        ("UC06", "Chặn người dùng", "Người dùng", "Ngăn user khác nhắn tin/gọi/xem trạng thái.", "Target tồn tại.", "BlockUser được lưu.", "1. Chọn chặn.\n2. Backend lưu BlockUser.\n3. Phát user-blocked.\n4. Middleware từ chối message/call liên quan.", "Unblock xóa BlockUser và phát user-unblocked."),
        ("UC07", "Tạo hội thoại nhóm", "Người dùng", "Tạo không gian chat nhiều người.", "Thành viên hợp lệ.", "Conversation type group được tạo.", "1. Chọn thành viên.\n2. Backend kiểm friendship/block.\n3. Lưu group name/admins/participants.\n4. Emit new-conversation.", "Nếu bật approval, thành viên mới vào queue chờ duyệt."),
        ("UC08", "Gửi tin nhắn", "Người dùng", "Lưu và phân phối tin realtime, sau đó hậu kiểm nội dung nếu cần.", "Người gửi thuộc hội thoại, không bị khóa.", "Message lưu và người nhận thấy realtime; moderation nền cập nhật trạng thái sau.", "1. Gửi content/media.\n2. Kiểm quyền.\n3. Upload media nếu có.\n4. Gắn moderationStatus = pending_review cho loại tin cần hậu kiểm.\n5. Lưu Message/Conversation.\n6. Emit new-message.\n7. Chạy moderation nền.\n8. Nếu vi phạm, emit message-moderated và ẩn nội dung.", "Nội dung vi phạm không bị chặn trước khi gửi mà bị đánh dấu/ẩn sau khi hậu kiểm xác nhận."),
        ("UC09", "Tương tác tin nhắn", "Người dùng", "Reply, reaction, recall, pin, forward.", "Message hợp lệ và user có quyền.", "Message cập nhật, realtime event phát.", "1. Chọn thao tác.\n2. Backend kiểm quyền.\n3. Cập nhật Message.\n4. Emit event tương ứng.", "Tin bị report/expired/recalled thì hạn chế thao tác."),
        ("UC10", "Tìm kiếm", "Người dùng", "Tìm user, conversation và message.", "Đã đăng nhập.", "Trả kết quả theo quyền.", "1. Nhập từ khóa.\n2. Backend normalize tiếng Việt.\n3. Query dữ liệu theo quyền.\n4. Trả kết quả hoặc NDJSON stream.", "Từ khóa rỗng hoặc quá ngắn trả kết quả rỗng."),
        ("UC11", "Disappearing messages", "Người dùng/Admin nhóm", "Bật chế độ tin tự xóa.", "Có quyền thay đổi setting.", "Tin mới tự hết hạn sau 24 giờ.", "1. Bật setting.\n2. Conversation lưu disableAt.\n3. Message mới có expiresAt.\n4. Worker soft-delete khi đến hạn.\n5. Emit dm:message-expired.", "Media chỉ xóa khi không còn message active tham chiếu."),
        ("UC12", "Gọi trực tiếp", "Người dùng", "Gọi audio/video 1-1.", "Hai user là bạn, không đang trong cuộc gọi khác.", "Hai bên nhận LiveKit token hoặc call kết thúc.", "1. Caller emit call-offer.\n2. Backend kiểm friend/block.\n3. Redis lưu call state.\n4. Receiver accept.\n5. Backend cấp token.\n6. Kết nối LiveKit.", "Receiver offline có thể nhận push; timeout trả no-answer."),
        ("UC13", "Group call", "Thành viên nhóm", "Gọi audio/video trong nhóm.", "Conversation là group active.", "Group call active và thành viên join/leave được đồng bộ.", "1. Emit group-call:start.\n2. Redis lưu state.\n3. Emit group-call:incoming.\n4. Thành viên join nhận token.\n5. Emit user-joined/user-left.", "Nhóm disbanded hoặc user không là thành viên thì trả error."),
        ("UC14", "Phòng họp LiveKit", "Host, participant", "Tạo/join phòng họp có waiting room.", "Host đăng nhập, room hợp lệ.", "Participant được admit hoặc reject.", "1. Host tạo meeting.\n2. Participant join room.\n3. Nếu cần duyệt, vào waitingRoom.\n4. Host admit/reject.\n5. Cấp token nếu admit.", "Waiting room timeout nếu host không xử lý."),
        ("UC15", "Tạo reminder", "Người dùng", "Nhắc hẹn cá nhân hoặc chung.", "Thời gian hợp lệ.", "Reminder pending được worker xử lý.", "1. Nhập content/remindAt/repeat.\n2. Lưu Reminder.\n3. BullMQ lên lịch.\n4. Đến hạn gửi notification.\n5. User snooze/dismiss.", "Repeat tính lần tiếp theo; shared reminder cho phép join/leave."),
        ("UC16", "Report vi phạm", "Người dùng", "Báo cáo message hoặc user.", "Target tồn tại, reporter có quyền.", "Report pending được tạo.", "1. Chọn report.\n2. Nhập reason.\n3. Backend lấy snapshot.\n4. Lưu Report.\n5. Admin thấy trong dashboard.", "Không cho report chính mình hoặc message không thuộc hội thoại."),
        ("UC17", "AI moderation", "Hệ thống, AI service", "Phát hiện nội dung vi phạm tiêu chuẩn cộng đồng bằng hậu kiểm.", "Message đã được lưu ở pending_review hoặc có report cần đánh giá.", "Message được cập nhật approved/rejected/skipped; report/admin review có thêm dữ liệu AI.", "1. Local signal phân tích nhanh.\n2. Gọi Gemini nếu cần.\n3. Voice dùng AssemblyAI transcript.\n4. So sánh confidence.\n5. Nếu an toàn, cập nhật approved/skipped và gửi notification.\n6. Nếu vi phạm, set reportStatus, cleanup media nếu cần và emit message-moderated.", "AI lỗi hoặc thiếu cấu hình thì ghi skipped, không chặn cứng khi thiếu bằng chứng."),
        ("UC18", "Admin review report", "Admin", "Xử lý report và ghi nhận vi phạm.", "Admin đã đăng nhập.", "Report resolved/dismissed; violation history cập nhật.", "1. Admin xem report.\n2. Có thể chạy AI review.\n3. Resolve violation/no_violation.\n4. Violation service tăng counter.\n5. Gửi notification/lock nếu cần.", "Admin có thể mở lại hoặc xem lịch sử resolved report theo user."),
        ("UC19", "Khóa/mở khóa tài khoản", "Admin, hệ thống", "Hạn chế user vi phạm.", "Có lý do hoặc vượt ngưỡng violation.", "User bị khóa/mở khóa và session cập nhật.", "1. Admin/system gọi lock.\n2. Set lock.isLocked.\n3. Xóa Session.\n4. Disconnect socket.\n5. Unlock reset counter nếu chọn.", "User bị khóa có thể gửi appeal."),
        ("UC20", "Kháng cáo tài khoản bị khóa", "Người dùng bị khóa, Admin", "Yêu cầu xem xét lại quyết định khóa.", "Tài khoản đang bị khóa.", "Appeal pending/approved/rejected.", "1. User gửi reason từ màn hình đăng nhập.\n2. Backend lưu LockAppeal.\n3. Admin review.\n4. Approve thì unlock; reject thì ghi note.", "Nếu đã có appeal pending thì không tạo trùng."),
        ("UC21", "Xem dashboard và observability", "Admin", "Giám sát hoạt động hệ thống.", "Admin đã đăng nhập.", "Hiển thị số liệu và audit log.", "1. Admin mở dashboard.\n2. Backend tổng hợp stats/metrics.\n3. Trả user, report, request, latency, CPU/memory.\n4. Admin drill-down theo user.", "Thiếu nguồn metrics thì hiển thị phần còn lại."),
    ]
    for uc in use_cases:
        add_use_case(doc, *uc)

    add_heading(doc, "4.4. Activity Diagram", 2)
    activity_rows = [
        ["Gửi tin nhắn", "Người dùng soạn nội dung -> Frontend gọi API -> Backend kiểm quyền/upload -> Lưu Message pending_review nếu cần -> Emit Socket.IO -> Moderation nền -> Approved/skipped thì gửi notification, rejected thì emit message-moderated."],
        ["Gọi trực tiếp", "Caller emit call-offer -> Backend kiểm friend/block/lock -> Redis lưu call state -> Receiver accept -> Cấp LiveKit token -> Kết nối room -> Kết thúc và ghi system message."],
        ["Report và xử lý vi phạm", "Reporter tạo report -> Lưu snapshot -> Admin review/AI review -> Resolve -> Violation service tăng counter -> Cảnh báo hoặc lock -> User appeal nếu bị khóa."],
        ["Reminder", "User tạo reminder -> Lưu DB -> BullMQ lên lịch -> Worker đến hạn -> Socket/Push/Email -> User snooze/dismiss/delete."],
        ["Disappearing", "Bật setting -> Message mới có expiresAt -> Worker sweep -> Soft-delete -> Cleanup media nếu cần -> Emit expired event."],
    ]
    add_table(doc, ["Luồng hoạt động", "Mô tả tuần tự"], activity_rows, widths=[Inches(1.7), Inches(5.8)], font_size=9)

    add_heading(doc, "4.5. Sequence Diagram", 2)
    add_sequence(doc, "4.5.1. Sequence Đăng nhập", [
        ["1", "User", "Frontend", "Nhập email/mật khẩu hoặc chọn Google."],
        ["2", "Frontend", "Auth API", "POST /api/auth/signin hoặc /api/auth/google/mobile."],
        ["3", "Auth API", "MongoDB User", "Kiểm tra user, password, role và lock status."],
        ["4", "Auth API", "MongoDB Session", "Tạo refresh token/session."],
        ["5", "Auth API", "Frontend", "Trả user profile và token."],
        ["6", "Frontend", "Socket.IO", "Kết nối realtime với token."],
        ["7", "Socket middleware", "MongoDB/Redis", "Xác thực session và join user room."],
    ])
    add_sequence(doc, "4.5.2. Sequence Gửi tin nhắn và hậu kiểm AI", [
        ["1", "Sender", "Frontend", "Gửi text/link/image/audio/file."],
        ["2", "Frontend", "Message API", "POST /api/messages/send."],
        ["3", "Message API", "Middleware", "Kiểm tra membership, block, lock, file."],
        ["4", "Message API", "MongoDB", "Lưu Message với moderationStatus = pending_review nếu loại tin cần hậu kiểm, cập nhật Conversation/unread."],
        ["5", "Message API", "Socket.IO", "Emit new-message ngay cho conversation room."],
        ["6", "Message API", "Background moderation", "setImmediate chạy moderateDeliveredMessage."],
        ["7", "Background moderation", "Gemini/AssemblyAI", "Local signal, Gemini và AssemblyAI transcript nếu là audio."],
        ["8", "Background moderation", "MongoDB/Socket.IO", "Nếu approved/skipped, cập nhật metadata và gửi notification; nếu rejected, set reportStatus, cleanup media và emit message-moderated."],
    ])
    add_sequence(doc, "4.5.3. Sequence Direct Call", [
        ["1", "Caller", "Socket.IO", "Emit call-offer(toUserId, callType)."],
        ["2", "Call handler", "MongoDB", "Kiểm tra receiver, friendship, block, lock."],
        ["3", "Call handler", "Redis", "Đặt offer lock và lưu call state."],
        ["4", "Call handler", "Receiver", "Emit incoming-call."],
        ["5", "Receiver", "Socket.IO", "Emit accept-call hoặc reject."],
        ["6", "Call handler", "LiveKit", "Cấp token room."],
        ["7", "Call handler", "Caller/Receiver", "Emit call-answered/call-accepted."],
        ["8", "Client", "LiveKit", "Truyền audio/video."],
        ["9", "Client", "Socket.IO", "Emit call-ended/cancelled; backend finalize state."],
    ])
    add_sequence(doc, "4.5.4. Sequence Report và Admin Review", [
        ["1", "Reporter", "Report API", "POST /api/reports/messages/:messageId hoặc /users/:userId."],
        ["2", "Report API", "MongoDB", "Kiểm quyền, lấy snapshot message/user."],
        ["3", "Report API", "MongoDB Report", "Tạo report pending."],
        ["4", "Admin", "Admin API", "GET /api/admin/reports."],
        ["5", "Admin", "AI review", "POST /api/admin/reports/messages/ai-review nếu cần."],
        ["6", "Admin", "Admin API", "PATCH /reports/:id/resolve."],
        ["7", "Admin API", "Violation service", "Ghi violation, tăng counter, gửi notification/lock."],
    ])
    add_sequence(doc, "4.5.5. Sequence Disappearing Message", [
        ["1", "User/Admin", "DM API", "PUT /api/dm/conversations/:id/disappearing."],
        ["2", "DM API", "MongoDB Conversation", "Lưu enabled, disableAt, enabledBy."],
        ["3", "DM API", "Socket.IO", "Emit dm:disappearing-setting-updated."],
        ["4", "Sender", "Message API", "Gửi message mới trong mode disappearing."],
        ["5", "Message API", "MongoDB Message", "Gắn expiresAt 24 giờ."],
        ["6", "Worker", "Internal DM API", "Expire batch định kỳ."],
        ["7", "Worker", "Socket.IO", "Emit dm:message-expired."],
    ])

    add_heading(doc, "4.6. Class Diagram", 2)
    class_rows = [
        ["User", "email, password, displayName, avatarUrl, bio, phone, profileVisibility, googleId, role, lock, moderation, fcmTokens", "Đại diện tài khoản, hồ sơ, quyền, khóa tài khoản và vi phạm.", "1-n Session, 1-1 UserStatus, n-n Conversation."],
        ["Session", "userId, refreshToken, expiresAt, deviceInfo, fcmTokens", "Quản lý phiên và thu hồi thiết bị.", "n-1 User."],
        ["UserStatus", "userId, manual_status, status_mode, last_seen_at", "Lưu presence và last seen.", "1-1 User."],
        ["Friend", "userA, userB, nicknameA, nicknameB", "Quan hệ bạn bè hai chiều.", "n-1 User ở hai đầu."],
        ["FriendRequest", "from, to, message, status", "Lời mời kết bạn.", "n-1 User gửi/nhận."],
        ["BlockUser", "from, to", "Quan hệ chặn một chiều.", "n-1 User."],
        ["Conversation", "type, directKey, participants, group, lastMessage, unreadCounts, disappearing, cleanup", "Hội thoại direct/group và cấu hình nhóm.", "1-n Message, n-n User."],
        ["Message", "conversationId, senderId, type, content, metadata, searchContent, filePublicId, replyTo, reactions, deliveredTo, expiresAt, reportStatus", "Tin nhắn, media, reaction, moderation và tự xóa.", "n-1 Conversation, n-1 User, self replyTo."],
        ["Meeting", "roomName, hostId, conversationId, status, participants, waitingRoom, requireApproval", "Phòng họp LiveKit.", "n-1 User host, n-1 Conversation tùy chọn."],
        ["Reminder", "userId, scope, sharedKey, conversationId, meetingId, content, remindAt, repeatRule, status, notifyChannels", "Nhắc hẹn cá nhân/chung.", "n-1 User, Conversation, Meeting."],
        ["Notification", "userId, title, content, linkUrl, type, actorId, metadata, isRead", "Thông báo trong app.", "n-1 User."],
        ["Report", "reporterId, targetType, targetUserId, targetMessageId, reasonCategory, status, snapshots, review, resolution", "Báo cáo vi phạm và kết quả review.", "n-1 User, 0..1 Message."],
        ["LockAppeal", "userId, email, reason, status, reviewedBy, adminNote", "Kháng cáo tài khoản bị khóa.", "n-1 User/Admin."],
        ["AuditLog", "userId, role, method, path, statusCode, durationMs, ip, userAgent", "Log request phục vụ quản trị.", "n-1 User."],
    ]
    add_table(doc, ["Lớp", "Thuộc tính chính", "Trách nhiệm", "Quan hệ"], class_rows, widths=[Inches(1.2), Inches(2.7), Inches(2.1), Inches(1.6)], font_size=8)

    add_heading(doc, "4.7. Data Model Diagram / ERD", 2)
    doc.add_picture(str(diagrams["data_model"]), width=Inches(6.7))
    add_caption(doc, "Hình 4.4. Data Model / ERD rút gọn")
    add_table(
        doc,
        ["Collection", "Khóa tham chiếu", "Trường chính", "Index / ràng buộc"],
        [
            ["users", "_id", "email, password, displayName, role, lock, moderation", "email unique; role/lock/createdAt index"],
            ["sessions", "userId -> users", "refreshToken, expiresAt, deviceInfo", "refreshToken unique; expiresAt TTL"],
            ["userstatuses", "userId -> users", "manual_status, status_mode, last_seen_at", "userId unique"],
            ["friends", "userA/userB -> users", "nicknameA, nicknameB", "unique userA + userB"],
            ["friendrequests", "from/to -> users", "message, status", "unique from + to; createdAt TTL"],
            ["blockusers", "from/to -> users", "createdAt", "unique from + to"],
            ["conversations", "participants.userId -> users", "type, directKey, group, lastMessage, unreadCounts, disappearing", "directKey unique cho direct; participants index"],
            ["messages", "conversationId -> conversations, senderId -> users, replyTo -> messages", "type, content, metadata, filePublicId, mentions, reactions, expiresAt, reportStatus", "conversationId + createdAt; mentions.userId; expiresAt"],
            ["meetings", "hostId -> users, conversationId -> conversations", "roomName, status, participants, waitingRoom", "roomName unique; ended meeting TTL"],
            ["reminders", "userId/createdBy -> users, conversationId, meetingId", "scope, sharedKey, content, remindAt, repeatRule, status", "sharedKey + userId unique; TTL theo status"],
            ["notifications", "userId -> users", "title, content, linkUrl, type, metadata, isRead", "userId + createdAt; createdAt TTL"],
            ["reports", "reporterId/targetUserId -> users, targetMessageId -> messages", "targetType, reasonCategory, snapshots, review, resolution", "status + createdAt; expiresAt TTL"],
            ["lockappeals", "userId/reviewedBy -> users", "email, reason, status, adminNote", "status + createdAt; expiresAt TTL"],
            ["auditlogs", "userId -> users", "method, path, statusCode, durationMs", "userId + createdAt; TTL 60 ngày"],
            ["pushsubscriptions", "userId -> users", "endpoint, keys, userAgent", "endpoint unique; TTL 90 ngày"],
            ["otps", "email logic -> users.email", "otp, type, expiresAt", "expiresAt TTL"],
        ],
        widths=[Inches(1.35), Inches(2.0), Inches(2.5), Inches(1.9)],
        font_size=8,
    )

    add_heading(doc, "4.8. Pipeline AI Moderation", 2)
    doc.add_picture(str(diagrams["pipeline"]), width=Inches(6.7))
    add_caption(doc, "Hình 4.3. Pipeline AI kiểm duyệt nội dung")
    add_para(doc, "Pipeline moderation của NexCon là pipeline hậu kiểm. Sau khi backend kiểm quyền, message được lưu và emit realtime trước với trạng thái pending_review. Tác vụ nền sau đó dùng local signal, Gemini và AssemblyAI để đánh giá nội dung. Nếu kết quả an toàn hoặc skipped, message được cập nhật metadata và notification mới được gửi; nếu vi phạm, message bị set reportStatus, nội dung UI được thay bằng thông báo vi phạm, media được cleanup nếu cần và conversation nhận event message-moderated.")
    add_table(
        doc,
        ["Bước", "Mô tả", "Đầu ra"],
        [
            ["1. Nhận và lưu", "Message API nhận text, link, image, file hoặc audio; kiểm quyền rồi lưu message pending_review nếu cần.", "Message đã emit realtime."],
            ["2. Chạy nền", "setImmediate gọi moderateDeliveredMessage sau khi message đã được lưu.", "Job hậu kiểm không chặn response gửi tin."],
            ["3. AI analysis", "Local signal, Gemini cho text/link/image/file metadata; AssemblyAI cho audio transcript nếu cần.", "blocked, category, confidence, reason."],
            ["4. Decision", "So sánh confidence với ngưỡng và cập nhật approved/rejected/skipped.", "approved/skipped hoặc rejected."],
            ["5. Hậu xử lý", "Approved/skipped thì gửi notification; rejected thì set reportStatus, cleanup media và emit message-moderated.", "Dữ liệu cập nhật và realtime event."],
        ],
        widths=[Inches(1.3), Inches(4.0), Inches(2.2)],
        font_size=9,
    )

    add_heading(doc, "4.9. API, realtime event và worker", 2)
    add_table(
        doc,
        ["Nhóm", "Endpoint / event / worker", "Chức năng"],
        [
            ["REST Auth", "/api/auth, /api/otp", "Signin, signup, OTP, sessions, reset password, Google, appeal."],
            ["REST User/Friend", "/api/users, /api/friends", "Profile, status, search, friend request, block, nickname."],
            ["REST Chat", "/api/conversations, /api/messages, /api/search, /api/dm", "Conversation, message, media, search, disappearing."],
            ["REST Call/Meeting", "/api/livekit, /api/meetings", "LiveKit token, meeting create/join/end."],
            ["REST Admin", "/api/admin, /api/reports", "Stats, observability, reports, users, appeals, locks."],
            ["Socket message", "new-message, read-message, typing, message-reaction, pin-message", "Đồng bộ chat realtime."],
            ["Socket call", "call-offer, incoming-call, accept-call, group-call:*", "Signaling direct/group call."],
            ["Socket system", "session-revoked, new-notification, reminder-triggered, dm:*", "Phiên, thông báo, reminder, disappearing."],
            ["Workers", "reminder, realtimeTimeout, disappearingMessage, groupCleanup, conversationClearCleanup", "Tác vụ nền và cleanup."],
        ],
        widths=[Inches(1.4), Inches(3.0), Inches(3.1)],
        font_size=8,
    )
    doc.add_page_break()


def add_chapter_5(doc):
    add_heading(doc, "CHƯƠNG 5. CÀI ĐẶT VÀ KIỂM THỬ", 1)
    add_heading(doc, "5.1. Cấu trúc mã nguồn", 2)
    add_table(
        doc,
        ["Thư mục", "Mô tả"],
        [
            ["backend/src/controllers", "Xử lý request cho auth, user, friend, conversation, message, meeting, reminder, report, admin."],
            ["backend/src/routes", "Định nghĩa REST endpoints."],
            ["backend/src/models", "Mongoose schema cho dữ liệu chính."],
            ["backend/src/socket", "Socket.IO gateway, direct call handler, group call handler."],
            ["backend/src/services", "Moderation, call state, notification, push, system metrics, LiveKit/Gemini helpers."],
            ["backend/src/workers", "BullMQ workers cho reminder, disappearing expiry, cleanup và timeout."],
            ["frontend/src/pages", "Các màn hình route chính của web app."],
            ["frontend/src/components", "UI components cho chat, call, admin, auth, reminder, people."],
            ["frontend/src/stores", "Zustand stores cho auth, chat, socket, call, meeting, reminder."],
            ["frontend/src/services", "Axios service gọi API backend."],
            ["docs", "Tài liệu kỹ thuật và báo cáo."],
        ],
        widths=[Inches(2.4), Inches(5.1)],
        font_size=9,
    )

    add_heading(doc, "5.2. Cài đặt môi trường phát triển", 2)
    add_table(
        doc,
        ["Thành phần", "Yêu cầu / ghi chú"],
        [
            ["Node.js", "Dùng để chạy frontend, backend và test."],
            ["MongoDB", "Database chính."],
            ["Redis", "Socket.IO adapter, call state, presence, BullMQ, violation counter."],
            ["Cloudinary", "Cấu hình cloud name, api key, api secret để upload media."],
            ["LiveKit", "Cấu hình URL, API key/secret để cấp token call/meeting."],
            ["Gemini/AssemblyAI", "Cấu hình API key nếu dùng AI moderation và transcription."],
            ["Firebase/Web Push", "Cấu hình FCM và VAPID cho notification."],
            ["Capacitor Android", "Dùng khi build/chạy Android app."],
        ],
        widths=[Inches(2.0), Inches(5.5)],
        font_size=9,
    )
    add_heading(doc, "5.3. Quy trình cài đặt tham khảo", 2)
    add_code(doc, """
cd backend
npm install
npm run dev

cd frontend
npm install
npm run dev
""")
    add_para(doc, "Khi chạy production hoặc staging cần cấu hình biến môi trường cho CLIENT_URL, FRONTEND_URL, JWT secret, MongoDB URI, Redis URL, Cloudinary, LiveKit, Gemini, AssemblyAI, Email, VAPID và Firebase.")

    add_heading(doc, "5.4. Deploy tích hợp CI/CD", 2)
    add_para(doc, "Quy trình CI/CD của NexCon được thiết kế để mọi thay đổi đi qua Pull Request vào nhánh main, chạy kiểm tra tự động bằng GitHub Actions, sau đó mới triển khai lên môi trường production. Frontend được deploy trên Vercel, backend được deploy trên Railway, còn MongoDB, Redis, LiveKit, Cloudinary, Firebase và các dịch vụ AI được cấu hình bằng biến môi trường.")
    add_table(
        doc,
        ["Thành phần", "Cấu hình / luồng xử lý"],
        [
            ["Repository", "Source code quản lý trên GitHub, nhánh production chính là main."],
            ["CI", "GitHub Actions tại .github/workflows/ci.yml chạy khi push main hoặc Pull Request vào main."],
            ["Backend check", "npm ci và npm test trong thư mục backend; test bằng Node.js test runner."],
            ["Frontend check", "npm ci --legacy-peer-deps, npm run lint, npm test và npm run build trong thư mục frontend."],
            ["Branch protection", "Khuyến nghị bật ruleset cho main: bắt buộc Pull Request, required checks Backend/Frontend pass, block force push."],
            ["CD backend", "Railway tự deploy backend sau khi main thay đổi; Root Directory: backend; Install Command: npm ci; Start Command: npm start; Health check: /api/auth/health."],
            ["CD frontend", "Vercel tự deploy frontend sau khi main thay đổi; Root Directory: frontend; Build Command: npm run build; Output Directory: dist."],
            ["Secrets", "Biến môi trường production được cấu hình trong Railway/Vercel, không commit .env, service account, private key hoặc keystore."],
        ],
        widths=[Inches(1.8), Inches(5.7)],
        font_size=9,
    )
    add_caption(doc, "Bảng 5.1. Cấu hình CI/CD và deploy production")
    add_code(doc, """
Feature branch
  -> Pull Request vào main
  -> GitHub Actions chạy Backend và Frontend checks
  -> Checks pass mới merge
  -> Merge main
  -> Railway deploy backend
  -> Vercel deploy frontend
  -> Theo dõi health check, logs và admin observability
""")
    add_para(doc, "Các biến môi trường tối thiểu của frontend gồm VITE_API_URL, VITE_SOCKET_URL, VITE_LIVEKIT_URL, VITE_VAPID_PUBLIC_KEY và VITE_CONNECTIVITY_CHECK_URL. Backend cần dùng cùng MongoDB connection string, Redis URL, JWT secrets, MESSAGE_ENCRYPTION_KEY, Cloudinary credentials, LiveKit key/secret, Firebase/Web Push và các API key AI đang bật.")

    add_heading(doc, "5.5. Xử lý backend multi-replica với 6 replicas", 2)
    add_para(doc, "Ở production, backend NexCon có thể chạy 6 replicas API phía sau load balancer để tăng khả năng chịu tải và giảm rủi ro một instance lỗi làm gián đoạn toàn bộ hệ thống. Vì hệ thống có Socket.IO, call state, presence, worker và moderation nền, các replica phải dùng chung Redis và MongoDB để giữ trạng thái nhất quán.")
    add_code(doc, """
Client Web/Android
  -> Load Balancer
    -> Backend Replica 1
    -> Backend Replica 2
    -> Backend Replica 3
    -> Backend Replica 4
    -> Backend Replica 5
    -> Backend Replica 6

Tất cả replicas
  -> MongoDB chung
  -> Redis chung: Socket.IO adapter, presence, call state, BullMQ, violation counter
  -> Cloudinary / LiveKit / FCM / Web Push / Gemini / AssemblyAI
""")
    add_table(
        doc,
        ["Vấn đề khi chạy 6 replicas", "Cách xử lý trong NexCon"],
        [
            ["Socket.IO event có thể phát từ bất kỳ replica nào", "Dùng @socket.io/redis-adapter để broadcast qua Redis pub/sub; socket join user:<userId> và conversationId room để nhận event xuyên replica."],
            ["Presence online/offline không nằm cố định ở một process", "Lưu presence, last seen và socket relation qua Redis/service để replica nào cũng đọc được trạng thái mới nhất."],
            ["Direct call/group call có thể được thao tác từ nhiều replica", "Lưu call state và lock trong Redis; offer/start/finalize lock giúp tránh double call, double accept hoặc nhiều thiết bị xử lý trùng."],
            ["Unread, delivered, read và message state cần nhất quán", "Lưu dữ liệu chính trong MongoDB, cập nhật Conversation/Message bằng index và kiểm tra membership trên mỗi request."],
            ["Moderation nền chạy sau khi message đã emit", "Replica nhận request sẽ schedule background moderation; kết quả cập nhật MongoDB và emit message-moderated qua Socket.IO Redis adapter để mọi client đều nhận."],
            ["Worker/scheduler có nguy cơ chạy trùng nếu bật ở mọi replica", "Khuyến nghị tách worker thành service riêng hoặc tắt inline worker trên API replicas bằng env như ENABLE_INLINE_DISAPPEARING_MESSAGE_WORKER=false, ENABLE_INLINE_GROUP_CLEANUP_WORKER=false, ENABLE_INLINE_CONVERSATION_CLEAR_CLEANUP_WORKER=false; các worker dùng chung Redis/BullMQ."],
            ["Long-polling qua load balancer có thể mất session", "Ưu tiên WebSocket transport ổn định; nếu cho phép polling fallback thì cấu hình sticky session trên load balancer."],
            ["Secrets khác nhau giữa replica gây lỗi token/mã hóa", "Tất cả 6 replicas phải dùng cùng ACCESS_TOKEN_SECRET, REFRESH_TOKEN_SECRET, MESSAGE_ENCRYPTION_KEY, MONGODB_CONNECTION_STRING và REDIS_URL."],
            ["Health check và rollout", "Dùng /api/auth/health cho Railway health check; rollout theo từng replica, theo dõi logs, error rate, latency và Redis connectivity."],
        ],
        widths=[Inches(2.55), Inches(4.95)],
        font_size=8,
    )
    add_caption(doc, "Bảng 5.2. Xử lý vận hành backend multi-replica 6 replicas")
    add_para(doc, "Mô hình khuyến nghị là 6 API replicas chỉ xử lý HTTP/Socket.IO request, còn reminder/disappearing/group cleanup/conversation cleanup workers chạy ở một hoặc vài worker service riêng. Cách này giảm rủi ro scheduler tạo job trùng, đồng thời vẫn cho phép BullMQ phân phối job an toàn qua Redis.")

    add_heading(doc, "5.6. Chiến lược kiểm thử", 2)
    add_bullets(doc, [
        "Unit test backend cho helper/middleware/service như moderation prompt, field format, mentions, disappearing messages và user status.",
        "Unit test frontend cho utility như meeting link, mentions, field format và disappearing messages.",
        "Integration test thủ công cho các luồng realtime do cần nhiều client/socket và dịch vụ ngoài.",
        "Kiểm thử bảo mật cho auth, role, session revoke, locked account, block user và signed media URL.",
        "Kiểm thử worker cho reminder, disappearing expiry và cleanup dữ liệu.",
        "Kiểm thử mobile cho FCM, local notification, incoming call và screenshot bridge.",
    ])

    add_heading(doc, "5.7. Test cases và kết quả mong đợi", 2)
    add_test_table(doc, [
        ["TC01", "Đăng ký", "Email chưa tồn tại.", "Nhập thông tin, nhận OTP, nhập OTP đúng.", "User được tạo, có thể đăng nhập."],
        ["TC02", "Đăng nhập", "User hợp lệ.", "Nhập email/mật khẩu đúng.", "Session tạo, token trả về, socket kết nối."],
        ["TC03", "Đăng nhập tài khoản bị khóa", "lock.isLocked=true.", "Thử đăng nhập.", "HTTP 423/restriction, hiển thị lý do và appeal."],
        ["TC04", "Quên mật khẩu", "Email đã đăng ký.", "Gửi OTP reset, đổi mật khẩu.", "Mật khẩu mới hoạt động, OTP hết hạn."],
        ["TC05", "Thu hồi session", "Có nhiều session.", "Xóa một session.", "Thiết bị tương ứng nhận session-revoked."],
        ["TC06", "Cập nhật hồ sơ", "User đăng nhập.", "Sửa bio/avatar.", "User cập nhật, event profile-updated."],
        ["TC07", "Gửi lời mời kết bạn", "Hai user chưa là bạn.", "A gửi request cho B.", "FriendRequest pending, B nhận notification."],
        ["TC08", "Chấp nhận kết bạn", "Request pending.", "B accept.", "Friend tạo, hai phía cập nhật realtime."],
        ["TC09", "Chặn người dùng", "A và B tồn tại.", "A block B, B gửi message.", "Message/call bị từ chối."],
        ["TC10", "Tạo nhóm", "Thành viên hợp lệ.", "Tạo group.", "Conversation group tạo, member nhận new-conversation."],
        ["TC11", "Gửi text an toàn", "User thuộc hội thoại.", "Gửi text bình thường.", "Message lưu, new-message emit."],
        ["TC12", "Gửi text vi phạm", "AI cấu hình.", "Gửi nội dung vi phạm.", "Tin hiển thị realtime trước ở pending_review, sau đó bị message-moderated và nội dung bị ẩn khi hậu kiểm xác nhận."],
        ["TC13", "Gửi ảnh", "File ảnh hợp lệ.", "Upload ảnh.", "Cloudinary lưu, message có filePublicId."],
        ["TC14", "Gửi voice", "Audio hợp lệ.", "Gửi voice.", "Audio lưu và emit trước, transcript được hậu kiểm nếu AssemblyAI cấu hình."],
        ["TC15", "Reaction", "Message hợp lệ.", "Thả emoji.", "reactions cập nhật, event message-reaction."],
        ["TC16", "Pin/Unpin", "Message hợp lệ.", "Pin rồi unpin.", "isPinned cập nhật, event pin-message."],
        ["TC17", "Recall", "Sender có quyền.", "Thu hồi tin.", "isRecalled=true, UI ẩn nội dung."],
        ["TC18", "Forward", "Source message hợp lệ.", "Forward đến hội thoại khác.", "Message mới tạo ở đích."],
        ["TC19", "Mention", "Có member nhóm.", "Gửi @member.", "Mention count tăng, user_mentioned emit."],
        ["TC20", "Search message", "Có message phù hợp.", "Tìm từ khóa.", "Trả message theo quyền."],
        ["TC21", "Direct call", "Hai user là bạn.", "Caller gọi, receiver accept.", "Hai bên nhận LiveKit token."],
        ["TC22", "Call timeout", "Receiver không trả lời.", "Chờ hết timeout.", "Caller nhận no-answer, Redis state dọn."],
        ["TC23", "Group call", "Group active.", "Start call, member join.", "group-call:user-joined emit."],
        ["TC24", "Meeting waiting room", "Host bật approval.", "Participant join, host admit.", "Participant nhận token."],
        ["TC25", "Reminder cá nhân", "remindAt tương lai.", "Tạo reminder.", "Worker phát reminder-triggered đúng hạn."],
        ["TC26", "Shared reminder", "User thuộc hội thoại.", "Tạo từ message.", "sharedKey tạo, member join/leave được."],
        ["TC27", "Notification", "Có event message/reminder.", "Kiểm tra danh sách.", "new-notification và mark read hoạt động."],
        ["TC28", "Disappearing setting", "Có quyền.", "Bật disappearing.", "Setting cập nhật, socket emit dm setting."],
        ["TC29", "Disappearing expiry", "Message có expiresAt.", "Chờ worker expire.", "Message soft-delete, dm:message-expired emit."],
        ["TC30", "Screenshot bridge", "Android 14+, disappearing active.", "Chụp màn hình.", "dm:screenshot-detected emit."],
        ["TC31", "Report message", "Message thuộc hội thoại.", "Tạo report.", "Report pending có snapshot."],
        ["TC32", "Admin AI review", "Có report pending.", "Chạy AI review.", "resolution.aiModeration có kết quả."],
        ["TC33", "Resolve report violation", "Admin đăng nhập.", "Resolve violation.", "Violation counter tăng, notification gửi."],
        ["TC34", "Auto lock", "Counter đạt ngưỡng.", "Ghi thêm violation.", "User bị lock, session xóa, socket ngắt."],
        ["TC35", "Appeal", "User bị khóa.", "Gửi appeal, admin approve.", "Lock tắt, user đăng nhập lại được."],
        ["TC36", "Admin observability", "Admin đăng nhập.", "Mở dashboard.", "Hiển thị request/error/latency/runtime metrics."],
        ["TC37", "CI/CD Pull Request", "Có branch mới và GitHub Actions cấu hình.", "Tạo Pull Request vào main.", "Backend và Frontend checks chạy; chỉ merge khi required checks pass."],
        ["TC38", "Deploy production", "Merge main thành công.", "Theo dõi Railway và Vercel deployment.", "Backend/frontend deploy bản mới, health check /api/auth/health pass."],
        ["TC39", "Backend 6 replicas", "6 API replicas dùng cùng MongoDB/Redis/secrets.", "Kết nối nhiều client vào các replica khác nhau, gửi message/call/group event.", "Socket.IO Redis adapter đồng bộ event, presence/call state nhất quán, không mất message."],
        ["TC40", "Worker tách API replica", "API replicas tắt inline cleanup workers, worker service riêng dùng cùng Redis.", "Tạo reminder/disappearing message và chờ worker xử lý.", "Job xử lý đúng một lần theo queue/lock, API replicas vẫn nhận realtime event cập nhật."],
    ])
    doc.add_page_break()


def add_chapter_6(doc):
    add_heading(doc, "CHƯƠNG 6. THIẾT KẾ GIAO DIỆN", 1)
    add_para(doc, "Giao diện NexCon được tổ chức theo các route chính trong frontend. Thiết kế ưu tiên thao tác nhanh, bố cục quen thuộc với ứng dụng giao tiếp và khả năng đồng bộ realtime giữa nhiều thiết bị.")
    add_table(
        doc,
        ["Route", "Màn hình", "Chức năng chính"],
        [
            ["/signin", "Đăng nhập", "Đăng nhập thường/Google, hiển thị trạng thái tài khoản bị khóa và form appeal."],
            ["/signup", "Đăng ký", "Tạo tài khoản, kiểm tra field, chuyển sang OTP."],
            ["/otp, /otp-resetpass, /reset-password", "OTP và reset password", "Xác minh email và đặt lại mật khẩu."],
            ["/chat", "Chat app", "Danh sách hội thoại, cửa sổ chat, message input, media, call, group info, disappearing toggle."],
            ["/people", "Bạn bè", "Friends, requests, suggestions, groups, blocked users."],
            ["/meet", "Phòng họp", "Tạo/join meeting, waiting room, copy link, schedule reminder."],
            ["/reminder, /reminders", "Reminder", "Danh sách, lịch, tạo/sửa, snooze, dismiss, shared reminder."],
            ["/notification", "Thông báo", "Danh sách notification, mark read/unread/delete."],
            ["/reports/my", "Lịch sử report", "Xem report đã gửi và trạng thái xử lý."],
            ["/moderation", "Trạng thái moderation", "Xem violation summary, restriction, violation history."],
            ["/settings/sessions", "Phiên đăng nhập", "Xem và thu hồi session."],
            ["/admin/overview", "Admin overview", "Stats, danh sách user, lock/unlock, drawer chi tiết."],
            ["/admin/observability", "Admin observability", "Request count, error rate, latency, runtime metrics."],
            ["/admin/reports/messages", "Báo cáo tin nhắn", "Review report message, AI review, resolve."],
            ["/admin/reports/users", "Báo cáo người dùng", "Review report user và xử lý target."],
            ["/admin/appeals", "Kháng cáo", "Duyệt appeal tài khoản bị khóa."],
            ["/terms, /privacy, /community-standards", "Pháp lý", "Điều khoản, quyền riêng tư, tiêu chuẩn cộng đồng."],
        ],
        widths=[Inches(2.0), Inches(2.0), Inches(3.5)],
        font_size=9,
    )
    add_heading(doc, "6.1. Nguyên tắc thiết kế giao diện", 2)
    add_bullets(doc, [
        "Chat là màn hình trung tâm, truy cập nhanh đến hội thoại, media, search, call và group info.",
        "Thông báo, reminder và moderation status cần rõ trạng thái để người dùng không bỏ lỡ sự kiện quan trọng.",
        "Admin dashboard cần ưu tiên bảng dữ liệu, bộ lọc, trạng thái report và thao tác xử lý nhanh.",
        "Mobile cần hỗ trợ incoming call, local notification, FCM và điều hướng phù hợp màn hình nhỏ.",
    ])
    doc.add_page_break()


def add_chapter_7(doc):
    add_heading(doc, "CHƯƠNG 7. KẾT LUẬN", 1)
    add_heading(doc, "7.1. Kết quả đạt được", 2)
    add_bullets(doc, [
        "Hoàn thiện tài liệu phân tích thiết kế tổng hợp cho hệ thống NexCon dựa trên nhiều mẫu báo cáo.",
        "Mô tả rõ bối cảnh, mục tiêu, phạm vi, yêu cầu chức năng và phi chức năng.",
        "Xây dựng use case, activity, sequence, class diagram và data model/ERD ở mức phục vụ triển khai.",
        "Phân tích riêng pipeline AI moderation, report, violation, lock account và appeal.",
        "Tổng hợp kiến trúc công nghệ, API/realtime/worker và bộ test case chức năng.",
    ])
    add_heading(doc, "7.2. Ưu điểm", 2)
    add_bullets(doc, [
        "Hệ thống có đầy đủ chat/call/reminder/notification/report/admin thay vì chỉ dừng ở chat cơ bản.",
        "Tách rõ REST API, Socket.IO, LiveKit, Redis, MongoDB và worker nền.",
        "Cơ chế moderation kết hợp AI và admin review giúp cân bằng tự động hóa và kiểm soát thủ công.",
        "Data model có TTL/index phù hợp cho phiên, OTP, notification, report, reminder và message expiry.",
    ])
    add_heading(doc, "7.3. Hạn chế", 2)
    add_bullets(doc, [
        "Mã hóa hiện tại là mã hóa ở tầng lưu trữ backend, chưa phải end-to-end encryption hoàn chỉnh.",
        "AI moderation phụ thuộc API key và chất lượng mô hình, cần giám sát false positive/false negative.",
        "Một số luồng realtime cần kiểm thử tải với nhiều client và nhiều backend replica.",
        "Repository hiện chưa có target iOS nên các native bridge mới áp dụng cho Android.",
    ])
    add_heading(doc, "7.4. Hướng phát triển", 2)
    add_bullets(doc, [
        "Bổ sung end-to-end encryption cho hội thoại quan trọng.",
        "Hoàn thiện dashboard đánh giá chất lượng AI moderation và cơ chế human-in-the-loop.",
        "Thêm iOS target với FCM/APNs, incoming call và screenshot bridge tương đương Android.",
        "Tối ưu search bằng search engine chuyên dụng khi dữ liệu message tăng lớn.",
        "Bổ sung kiểm thử tự động end-to-end cho chat, call, admin report và worker.",
        "Thiết kế hệ thống audit/anomaly detection nâng cao để phát hiện spam hoặc abuse theo hành vi.",
    ])
    doc.add_page_break()


def add_references_and_appendix(doc):
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "Tài liệu README của dự án NexCon trong workspace c:\\HCMUTE\\NexCon.",
        "Tài liệu docs/dm-disappearing-messages.md của dự án NexCon.",
        "Mã nguồn backend: routes, controllers, models, services, socket handlers và workers.",
        "Mã nguồn frontend: routes, pages, stores, services và components.",
        "Các file DOCX mẫu được cung cấp: CLC_Nhom1_PhanThiMyLinh_NguyenHuuLoc_VuongTriHung.docx, ThamKhao.docx, Báo cáo PTTK HTTT.docx, Final Report.docx.",
        "Tài liệu chính thức của React, Vite, Node.js, Express, MongoDB, Redis, Socket.IO, LiveKit, Cloudinary, Firebase Cloud Messaging và BullMQ.",
    ]
    add_numbered(doc, refs)
    doc.add_page_break()

    add_heading(doc, "PHỤ LỤC", 1)
    add_heading(doc, "A. Đường dẫn liên kết đến đồ án", 2)
    add_table(
        doc,
        ["Mục", "Đường dẫn / thông tin"],
        [
            ["Repository", "https://github.com/ToiLaBao2004/NexCon"],
            ["Thư mục local", str(ROOT)],
            ["Frontend", str(ROOT / "frontend")],
            ["Backend", str(ROOT / "backend")],
            ["Tài liệu disappearing messages", str(ROOT / "docs" / "dm-disappearing-messages.md")],
            ["File báo cáo tạo mới", str(OUT)],
        ],
        widths=[Inches(2.0), Inches(5.5)],
        font_size=10,
    )
    add_heading(doc, "B. Gợi ý cập nhật sau khi mở bằng Microsoft Word", 2)
    add_bullets(doc, [
        "Điền thông tin giảng viên, lớp học phần và danh sách thành viên nhóm ở trang bìa.",
        "Nếu cần mục lục có số trang tự động, dùng References -> Table of Contents trong Word.",
        "Có thể thay thế hình minh họa bằng sơ đồ UML vẽ bằng StarUML, draw.io hoặc PlantUML nếu môn học yêu cầu hình UML chuẩn.",
        "Có thể chụp thêm giao diện thực tế của app và chèn vào Chương 6.",
    ])


def finalize_document(doc):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "NexCon - Báo cáo phân tích thiết kế hệ thống thông tin"
        for run in footer.runs:
            set_run_font(run, size=10)
    if OUT.exists():
        OUT.unlink()
    doc.save(str(OUT))


def main():
    diagrams = create_diagrams()
    doc = setup_document()
    add_front_matter(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc, diagrams)
    add_chapter_5(doc)
    add_chapter_6(doc)
    add_chapter_7(doc)
    add_references_and_appendix(doc)
    finalize_document(doc)
    print(OUT)


if __name__ == "__main__":
    main()
